import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from shapely.geometry import Polygon, box
import random
import pandas as pd
from fpdf import FPDF
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Dimensiunile utilajelor (Lungime x Lățime)
machines = [
    ("Platforma pentru lucru la inaltime - tip foarfeca 1", 1.66, 0.76, 6),
    ("Platforma pentru lucru la inaltime - tip foarfeca 2", 2.26, 1.16, 2),
    ("Platforma pentru lucru la inaltime", 2.26, 0.81, 2),
    ("Platforma pentru lucru la inaltime - tip foarfeca 3", 2.26, 0.81, 2),
    ("Platforma pentru lucru la inaltime cu brat articulat 1", 1.83, 0.76, 2),
    ("Platforma pentru lucru la inaltime cu brat articulat 2", 1.42, 0.76, 2),
    ("Platforma pentru lucru la inaltime cu brat articulat 3", 1.12, 0.70, 2),
    ("Rampa mobila", 1.83, 0.72, 1),
    ("Stalpi de iluminat fotovoltaici mobili", 0.114, 0.114, 4),
    ("Toaleta ecologica", 2.20, 1.60, 1),
    ("Drujba", 0.90, 0.25, 1),
    ("Tocator resturi vegetale", 0.707, 0.388, 1),
    ("Buldoexcavator", 3.40, 1.41, 1),
    ("Container de tip birou", 6.058, 2.438, 1),
    ("Generator", 3.00, 1.50, 1)
]

# Funcție pentru plottarea poligonului și utilajelor
def plot_polygon_with_machines(polygon_coords, placements):
    fig, ax = plt.subplots(figsize=(10, 8))

    # Plotăm poligonul
    polygon_coords.append(polygon_coords[0])
    x_coords, y_coords = zip(*polygon_coords)
    ax.plot(x_coords, y_coords, marker='o')
    polygon_coords.pop()  # Eliminăm punctul adăugat pentru închidere

    # Setăm aspectul graficului
    ax.set_aspect('equal', adjustable='box')
    ax.set_xticks([])
    ax.set_yticks([])
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # Plasăm utilajele
    for (x, y, w, h, label) in placements:
        rect = patches.Rectangle((x, y), w, h, linewidth=1, edgecolor='r', facecolor='none')
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha='center', va='center', fontsize=8, color='black')

    plt.xlabel('X')
    plt.ylabel('Y')
    plt.title('Aranjarea automată a utilajelor pe plot')
    plt.grid(True)
    return fig

# Funcție pentru a verifica dacă dreptunghiul este în interiorul poligonului
def is_rect_inside_polygon(rect, polygon):
    rect_box = box(rect[0], rect[1], rect[0] + rect[2], rect[1] + rect[3])
    return polygon.contains(rect_box)

# Funcție pentru a verifica suprapunerea dintre două dreptunghiuri
def is_overlapping(rect1, rect2):
    r1 = box(rect1[0], rect1[1], rect1[0] + rect1[2], rect1[1] + rect1[3])
    r2 = box(rect2[0], rect2[1], rect2[0] + rect2[2], rect2[1] + rect2[3])
    return r1.intersects(r2)

# Funcție pentru a plasa automat dreptunghiurile în interiorul poligonului
def place_machines_in_polygon(machines, polygon):
    placements = []
    legend = []
    label = 'A'
    for name, length, width, quantity in machines:
        added_to_legend = False
        for _ in range(quantity):
            placed = False
            while not placed:
                minx, miny, maxx, maxy = polygon.bounds
                x = random.uniform(minx, maxx - length)
                y = random.uniform(miny, maxy - width)
                rect = (x, y, length, width)
                if is_rect_inside_polygon(rect, polygon) and all(not is_overlapping(rect, p) for p in placements):
                    placements.append((x, y, length, width, label))
                    if not added_to_legend:
                        legend.append((label, name, quantity, f"{length} x {width}"))
                        added_to_legend = True
                    placed = True
        label = chr(ord(label) + 1)  # Incrementăm labelul
    return placements, legend

# Funcție pentru a genera PDF
def generate_pdf(fig, legend_df):
    buffer = BytesIO()
    fig.savefig(buffer, format='png')
    buffer.seek(0)
    image = buffer.read()

    pdf = FPDF()
    pdf.add_page()
    pdf.image(image, x=10, y=10, w=190)
    pdf.set_xy(10, 150)
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Legenda:", ln=True, align="L")

    for _, row in legend_df.iterrows():
        pdf.cell(200, 10, txt=f"{row['Identificator']}: {row['Denumire Utilaj']} - {row['Numar Bucati']} buc - {row['Dimensiune (L x l)']}", ln=True, align="L")

    return pdf.output(dest="S").encode("latin1")

# Funcție pentru a genera Word
def generate_word(fig, legend_df):
    buffer = BytesIO()
    fig.savefig(buffer, format='png')
    buffer.seek(0)
    image = buffer.read()

    doc = Document()
    doc.add_heading('Aranjarea automată a utilajelor pe plot', 0)

    # Adăugăm imaginea
    doc.add_picture(BytesIO(image), width=Inches(6))

    doc.add_heading('Legenda:', level=1)
    for _, row in legend_df.iterrows():
        doc.add_paragraph(f"{row['Identificator']}: {row['Denumire Utilaj']} - {row['Numar Bucati']} buc - {row['Dimensiune (L x l)']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Definim coordonatele poligonului
polygon_coords = [
    (399485.06, 385140.713),
    (399483.58, 385140.062),
    (399477.304, 385124.575),
    (399484.305, 385122.896),
    (399492.273, 385120.567),
    (399496.347, 385138.383)
]

polygon = Polygon(polygon_coords)

st.title("Aranjarea utilajelor pe plot")

# Plasăm automat utilajele în poligon
placements, legend = place_machines_in_polygon(machines, polygon)

# Plottăm poligonul și utilajele
fig = plot_polygon_with_machines(polygon_coords, placements)

# Creăm tabelul cu legenda
legend_df = pd.DataFrame(legend, columns=["Identificator", "Denumire Utilaj", "Numar Bucati", "Dimensiune (L x l)"])
st.table(legend_df)

# Adăugăm butoanele pentru descărcare PDF și Word
if st.button("Descarcă PDF"):
    pdf_data = generate_pdf(fig, legend_df)
    st.download_button(label="Descarcă PDF", data=pdf_data, file_name="utilaje.pdf", mime="application/pdf")

if st.button("Descarcă Word"):
    word_data = generate_word(fig, legend_df)
    st.download_button(label="Descarcă Word", data=word_data, file_name="utilaje.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

